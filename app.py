import streamlit as st
import openai
from google import genai
import os
import tempfile
import base64
import copy
import io
import re
import datetime
import time
from docx import Document
from docx.oxml.ns import qn

def extraer_vista_previa(docx_bytes):
    """Devuelve una lista de líneas de texto (y marcadores de imagen) para
    mostrar como vista previa, sin necesidad de descargar el archivo."""
    doc = Document(io.BytesIO(docx_bytes))
    lineas = []
    for el in doc.element.body:
        if el.tag == qn('w:p'):
            tiene_imagen = len(el.findall('.//' + qn('w:drawing'))) > 0
            texto = ''.join(
                t.text or '' for r in el.findall(qn('w:r')) for t in [r.find(qn('w:t'))] if t is not None
            ).strip()
            if tiene_imagen:
                lineas.append("🖼️ *(imagen incluida aquí)*")
            elif texto:
                lineas.append(texto)
    return lineas


CARPETA_HISTORIAL = "informes_generados"
os.makedirs(CARPETA_HISTORIAL, exist_ok=True)


def guardar_en_historial(nombre_archivo, contenido_bytes_o_texto):
    """Guarda una copia del informe generado en la carpeta de historial."""
    ruta = os.path.join(CARPETA_HISTORIAL, nombre_archivo)
    modo = "w" if isinstance(contenido_bytes_o_texto, str) else "wb"
    encoding = "utf-8" if modo == "w" else None
    with open(ruta, modo, encoding=encoding) as f:
        f.write(contenido_bytes_o_texto)


def _llenar_campo(tpl_body, etiqueta_texto, valor):
    """Escribe un valor justo después de una etiqueta con ':' en la plantilla
    (por ejemplo 'Cedula:', 'Medico Referente:', 'Fecha De Estudio:')."""
    if not valor:
        return
    for el in tpl_body:
        if el.tag == qn('w:p'):
            texto_parrafo = ''.join(el.itertext())
            if etiqueta_texto in texto_parrafo:
                runs = el.findall(qn('w:r'))
                run_colon = None
                for r in runs:
                    t = r.find(qn('w:t'))
                    if t is not None and t.text and ':' in t.text:
                        run_colon = r
                if run_colon is not None:
                    nuevo_run = copy.deepcopy(run_colon)
                    t_nuevo = nuevo_run.find(qn('w:t'))
                    t_nuevo.text = '  ' + str(valor)
                    t_nuevo.set(qn('xml:space'), 'preserve')
                    run_colon.addnext(nuevo_run)
                return


def unir_informe_renal(template_path, contenido_docx_file, fecha_estudio=None, medico_referente=None, cedula=None):
    """Une un documento Word del Dr. Quijada (hallazgos + imágenes) con la
    plantilla institucional de Gammagrama Renal. No usa IA: solo copia
    texto e imágenes, y detecta datos (nombre del paciente, M.S.A.S., C.M.)
    por patrones simples de texto."""
    tpl_doc = Document(template_path)
    content_doc = Document(contenido_docx_file)

    tpl_body = tpl_doc.element.body
    content_body = content_doc.element.body

    def encontrar_indice(body, condicion, desde=0):
        for i in range(desde, len(body)):
            el = body[i]
            if el.tag == qn('w:p') and condicion(''.join(el.itertext())):
                return i
        return None

    paciente_idx = encontrar_indice(content_body, lambda t: t.strip().startswith('Paciente:'))
    if paciente_idx is None:
        raise ValueError("No se encontró la línea 'Paciente:' en el documento del doctor.")
    atentamente_idx = encontrar_indice(content_body, lambda t: t.strip().startswith('Atentamente'), desde=paciente_idx + 1)
    if atentamente_idx is None:
        raise ValueError("No se encontró la línea 'Atentamente' (cierre/firma) en el documento del doctor.")

    insert_elements = [content_body[i] for i in range(paciente_idx + 1, atentamente_idx)]

    rid_map = {}
    for el in insert_elements:
        for blip in el.findall('.//' + qn('a:blip')):
            old_rid = blip.get(qn('r:embed'))
            if old_rid and old_rid not in rid_map:
                image_part = content_doc.part.related_parts[old_rid]
                new_rid, _ = tpl_doc.part.get_or_add_image(io.BytesIO(image_part.blob))
                rid_map[old_rid] = new_rid

    new_elements = []
    for el in insert_elements:
        new_el = copy.deepcopy(el)
        for blip in new_el.findall('.//' + qn('a:blip')):
            old_rid = blip.get(qn('r:embed'))
            if old_rid in rid_map:
                blip.set(qn('r:embed'), rid_map[old_rid])
        new_elements.append(new_el)

    title_idx = encontrar_indice(tpl_body, lambda t: 'ESTUDIO GAMMAGRAMA RENAL' in t)
    if title_idx is None:
        raise ValueError("No se encontró el título 'ESTUDIO GAMMAGRAMA RENAL' en la plantilla.")

    sig_idx = None
    for i in range(title_idx + 1, len(tpl_body)):
        el = tpl_body[i]
        if el.tag == qn('w:p') and el.findall('.//' + qn('w:drawing')):
            sig_idx = i
            break
    if sig_idx is None:
        raise ValueError("No se encontró la imagen de firma en la plantilla.")

    sig_el = tpl_body[sig_idx]

    for el in [tpl_body[i] for i in range(title_idx + 1, sig_idx)]:
        el.getparent().remove(el)

    for new_el in new_elements:
        sig_el.addprevious(new_el)

    nombre_paciente = None
    m_nombre = re.search(r'Paciente:\s*([^.:]+?)\.', ''.join(content_body[paciente_idx].itertext()))
    if m_nombre:
        nombre = m_nombre.group(1).strip()
        nombre_paciente = nombre
        for el in tpl_body:
            if el.tag == qn('w:p') and 'Paciente:' in ''.join(el.itertext()):
                for r in el.findall(qn('w:r')):
                    t = r.find(qn('w:t'))
                    if t is not None and t.text and t.text.strip() == '' and len(t.text) > 5:
                        t.text = ('  ' + nombre).ljust(len(t.text))
                        break
                break

    texto_completo = '\n'.join(''.join(el.itertext()) for el in content_body if el.tag == qn('w:p'))
    credenciales = []
    for patron in [r'M\.S\.A\.S\.?\s*\d+', r'C\.M\.?\s*\d+']:
        m = re.search(patron, texto_completo)
        if m and m.group(0) not in credenciales:
            credenciales.append(m.group(0))

    credential_style_p = None
    for el in tpl_body:
        if el.tag == qn('w:p') and 'Radioterapeuta' in ''.join(el.itertext()):
            credential_style_p = el

    if credential_style_p is not None:
        for cred in credenciales:
            nuevo_p = copy.deepcopy(credential_style_p)
            runs = nuevo_p.findall(qn('w:r'))
            for extra in runs[1:]:
                nuevo_p.remove(extra)
            t_el = runs[0].find(qn('w:t'))
            t_el.text = cred
            t_el.set(qn('xml:space'), 'preserve')
            credential_style_p.addnext(nuevo_p)
            credential_style_p = nuevo_p

    if fecha_estudio:
        _llenar_campo(tpl_body, 'Fecha De Estudio', fecha_estudio)
    if medico_referente:
        _llenar_campo(tpl_body, 'Medico Referente', medico_referente)
    if cedula:
        _llenar_campo(tpl_body, 'Cedula', cedula)

    buffer = io.BytesIO()
    tpl_doc.save(buffer)
    buffer.seek(0)
    return buffer, nombre_paciente


st.set_page_config(
    page_title="Medicina Nuclear - Sistema de Transcripción",
    page_icon="🏥",
    layout="wide"
)

with open("logo.png", "rb") as image_file:
    logo_base64 = base64.b64encode(image_file.read()).decode()

st.markdown(f"""
    <style>
    .block-container {{ padding-top: 3.5rem !important; }}
    </style>
    <div style="
        background-color: #ebe20e; 
        width: 100%; 
        padding: 15px 20px; 
        margin-bottom: 25px;
        box-shadow: 0px 4px 6px rgba(0,0,0,0.08);
        display: flex;
        align-items: center;
        border-radius: 5px;
    ">
        <img src="data:image/png;base64,{logo_base64}" width="150">
    </div>
""", unsafe_allow_html=True)

st.markdown("""
    <style>
    .main-title { font-size: 28px; font-weight: bold; color: #ebe20e; }
    .sub-title { font-size: 15px; color: #4B5563; margin-bottom: 25px; }
    .stButton button { background-color: #2563EB; color: white; border-radius: 8px; font-weight: bold; width: 100%; }
    </style>
""", unsafe_allow_html=True)

st.markdown('<div class="main-title"> Unidad de Medicina Nuclear</div>', unsafe_allow_html=True)
st.markdown('<div class="sub-title">Panel de Transcripción y Estructuración de Informes Médicos</div>', unsafe_allow_html=True)

st.sidebar.header("🔑 Claves de API")
openai_key = st.sidebar.text_input("OpenAI API Key", type="password", value=os.getenv("OPENAI_API_KEY", ""))
gemini_key = st.sidebar.text_input("Gemini API Key", type="password", value=os.getenv("GEMINI_API_KEY", ""))

tipo_estudio = st.sidebar.selectbox("Selecciona el estudio:", ["Gammagrafía Ósea", "Gammagrafía Tiroidea", "Rastreo Corporal Total", "Gammagrafía Renal (DTPA/DMSA)", "Plantilla Libre"])

PLANTILLAS = {
    "Gammagrafía Ósea": """I. DATOS TÉCNICOS

- Estudio: Rastreo corporal óseo
- Radiofármaco: MDP (Metilendifosfonato) marcado con Tc-99m (Tecnecio)
- Actividad administrada: 20 mCi
- Vía de administración: Endovenosa
- Equipo: Gammacámara Elscint Apex 409 AG
- Proyecciones obtenidas: Anterior y posterior de cuerpo entero
- Calidad técnica del estudio: Sin obstrucciones


II. ANTECEDENTES CLÍNICOS
Historia clínica, gammagramas previos, motivo de estudio


III. HALLAZGOS
Se realizó gammagrafía ósea total, observándose la distribución del radiofármaco en las siguientes regiones:

Vista anterior (proyección ANT)
- Cráneo y macizo facial:
- Escápulas y esternón:
- Parrilla costal anterior:
- Miembros superiores:
- Pelvis:
- Miembros inferiores:

Vista posterior (proyección POST)
- Columna vertebral:
- Pelvis posterior:
- Parrilla costal posterior y escápulas:
- Cráneo posterior:
- Sistema renal y tejidos blandos:


IV. CONCLUSIÓN DIAGNÓSTICO
""",
    "Gammagrafía Tiroidea": "INFORMACIÓN DEL ESTUDIO: GAMMAGRAFÍA TIROIDEA\n[DATOS DEL PACIENTE]\nINDICACIÓN:\n\nHALLAZGOS:\n- Glándula tiroides de morfología y ubicación:\n- Captación del trazador:\n- Nódulos / Lesiones:\n\nCONCLUSIÓN:",
    "Rastreo Corporal Total": "INFORMACIÓN DEL ESTUDIO: RASTREO CORPORAL TOTAL\n[DATOS DEL PACIENTE]\nINDICACIÓN:\n\nHALLAZGOS:\n- Áreas de hipercaptación fisiológica y patológica:\n\nCONCLUSIÓN:",
    "Gammagrafía Renal (DTPA/DMSA)": "INFORMACIÓN DEL ESTUDIO: GAMMAGRAFÍA RENAL\n[DATOS DEL PACIENTE]\nINDICACIÓN:\n\nHALLAZGOS:\n- Perfusión y función renal izquierda:\n- Perfusión y función renal derecha:\n- Excreción / Función relativa:\n\nCONCLUSIÓN:",
    "Plantilla Libre": "Insertar informe médico estructurado."
}

INSTRUCCIONES_ESTRICTAS_OSEA = """
REGLAS ADICIONALES OBLIGATORIAS PARA ESTE INFORME (Gammagrafía Ósea):
1. Respeta EXACTAMENTE la estructura de la plantilla: I. DATOS TÉCNICOS, II. ANTECEDENTES CLÍNICOS, III. HALLAZGOS (con sus dos subsecciones "Vista anterior (proyección ANT)" y "Vista posterior (proyección POST)", cada una con sus mismos ítems en el mismo orden), y IV. CONCLUSIÓN DIAGNÓSTICO.
2. NO fusiones, elimines, renombres ni reordenes ninguna subsección o ítem de "Vista anterior" ni de "Vista posterior", aunque el dictado los mencione en otro orden o de forma mezclada. Coloca cada hallazgo dictado en el ítem anatómico que le corresponda.
3. En "I. DATOS TÉCNICOS", los valores ya vienen fijos en la plantilla (radiofármaco, actividad, vía, equipo, proyecciones, calidad técnica). Mantenlos tal cual salvo que el dictado mencione explícitamente un valor distinto para ese campo puntual; en ese caso, usa el valor del dictado solo para ese campo.
4. Si un ítem de "Vista anterior" o "Vista posterior" no fue mencionado en el dictado, escribe exactamente "Dentro de límites normales" en ese ítem. No lo dejes vacío y no inventes hallazgos.
5. No agregues secciones, encabezados ni comentarios que no estén en la plantilla original.
"""

col1, col2 = st.columns([1, 1])

with col1:
    st.subheader("1. Cargar Dictado en Audio")
    st.caption("Puedes subir más de un audio si el médico envió varios para el mismo paciente (se combinan antes de generar el informe).")
    audio_files = st.file_uploader(
        "Arrastra el audio o los audios (.ogg, .mp3, .opus, .wav, .m4a)",
        type=["ogg", "mp3", "opus", "wav", "m4a"],
        accept_multiple_files=True
    )
    if audio_files:
        for i, af in enumerate(audio_files, start=1):
            st.write(f"Audio {i}: {af.name}")
            st.audio(af)
            if af.size < 20_000:
                st.warning(f"⚠️ El audio {i} ({af.name}) parece muy corto o vacío. Puede que la transcripción salga incompleta.")
    plantilla_actual = st.text_area("Plantilla a completar:", value=PLANTILLAS[tipo_estudio], height=320 if tipo_estudio == "Gammagrafía Ósea" else 220)

with col2:
    caja_resultados = st.container(border=True)
    with caja_resultados:
        st.subheader("2. Informe Final Generado")
        if not audio_files: st.info("Sube uno o más audios a la izquierda para comenzar.")

if audio_files and st.button("🚀 Procesar e Generar Informe"):
    if not openai_key or not gemini_key:
        st.error("⚠️ Ingrese ambas API Keys en la barra lateral.")
    else:
        # --- Barra de progreso general ---
        barra_progreso = st.progress(0, text="Iniciando...")

        transcripciones = []
        client_openai = openai.OpenAI(api_key=openai_key)
        total_audios = len(audio_files)

        for i, af in enumerate(audio_files, start=1):
            # La transcripción ocupa el primer 60% de la barra
            porcentaje = int((i - 1) / total_audios * 60)
            barra_progreso.progress(porcentaje, text=f"🎧 Transcribiendo audio {i} de {total_audios}...")

            ext = af.name.split('.')[-1].lower()
            if ext == 'opus': ext = 'ogg'
            with tempfile.NamedTemporaryFile(delete=False, suffix=f".{ext}") as tmp:
                tmp.write(af.read())
                tmp_path = tmp.name
            try:
                with open(tmp_path, "rb") as f:
                    transcript = client_openai.audio.transcriptions.create(model="whisper-1", file=f, language="es")
                transcripciones.append(transcript.text)
            except Exception as e:
                st.error(f"Error al transcribir el audio {i} ({af.name}): {e}")
            finally:
                os.remove(tmp_path)

        barra_progreso.progress(60, text="🎧 Transcripción completa.")

        if transcripciones:
            # Combina todos los audios en un solo texto, etiquetando cada parte
            texto_transcrito = "\n\n".join(
                f"[Audio {i}]\n{texto}" for i, texto in enumerate(transcripciones, start=1)
            )
            with st.expander("Ver texto crudo (todos los audios)"): st.write(texto_transcrito)

            try:
                client_gemini = genai.Client(api_key=gemini_key)
                instrucciones_extra = INSTRUCCIONES_ESTRICTAS_OSEA if tipo_estudio == "Gammagrafía Ósea" else ""
                nota_multi_audio = (
                    "\nNOTA: El dictado puede venir dividido en varios audios (marcados como [Audio 1], [Audio 2], etc.) "
                    "del mismo paciente. Combina la información de todos ellos en un solo informe coherente, "
                    "sin repetir datos duplicados ni mencionar que venían separados en audios distintos.\n"
                    if len(transcripciones) > 1 else ""
                )
                prompt = (
                    f"Eres un experto en medicina nuclear. Rellena la plantilla con el dictado. "
                    f"IMPORTANTE: no inventes ni infieras ningún dato que no esté explícitamente mencionado en el dictado. "
                    f"Si un campo de la plantilla no fue mencionado, escribe exactamente 'Dentro de límites normales' en ese campo, "
                    f"nunca lo completes con información supuesta. Usa negritas (Markdown **) para los títulos."
                    f"{nota_multi_audio}"
                    f"{instrucciones_extra}\n"
                    f"DICTADO: {texto_transcrito}\n"
                    f"PLANTILLA:\n{plantilla_actual}"
                )

                # --- Llamada a Gemini con reintentos automáticos ---
                MAX_INTENTOS = 4
                response = None
                ultimo_error = None
                for intento in range(MAX_INTENTOS):
                    # La generación con Gemini ocupa el 60%-95% de la barra
                    porcentaje = 60 + int((intento / MAX_INTENTOS) * 35)
                    if intento == 0:
                        barra_progreso.progress(porcentaje, text="🤖 Organizando informe con Gemini...")
                    else:
                        barra_progreso.progress(porcentaje, text=f"⏳ Servidor ocupado, reintentando ({intento + 1}/{MAX_INTENTOS})...")
                    try:
                        response = client_gemini.models.generate_content(
                            model="gemini-3.5-flash",
                            contents=prompt
                        )
                        break
                    except Exception as err_intento:
                        ultimo_error = err_intento
                        if "503" in str(err_intento) or "UNAVAILABLE" in str(err_intento):
                            if intento < MAX_INTENTOS - 1:
                                time.sleep(5 * (intento + 1))
                        else:
                            raise

                if response is None:
                    barra_progreso.progress(100, text="❌ No se pudo completar.")
                    raise ultimo_error

                barra_progreso.progress(95, text="💾 Guardando informe...")

                timestamp = datetime.datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
                nombre_hist = f"{timestamp}_{tipo_estudio.replace(' ', '_')}.txt"
                guardar_en_historial(nombre_hist, response.text)

                barra_progreso.progress(100, text="✅ ¡Listo!")

                with caja_resultados:
                    st.success("✨ ¡Informe listo! (se guardó una copia en el historial)")
                    st.markdown(response.text)
                    st.markdown("---")
                    st.text_area("Copiar para sistema:", value=response.text, height=200)
                    st.download_button("📥 Descargar Informe (.txt)", data=response.text, file_name="informe_medico.txt", mime="text/plain")
            except Exception as e:
                barra_progreso.progress(100, text="❌ Error.")
                st.error(f"Error al estructurar: {e}")

st.markdown("---")
st.header("📎 Unir Informe Renal (Dr. Quijada) con Plantilla")
st.markdown("Sube el documento Word que envía el Dr. Quijada (con sus hallazgos e imágenes) y la app lo une automáticamente con la plantilla institucional. No usa IA, por lo que no necesita las claves de API.")

doc_quijada = st.file_uploader("Documento del Dr. Quijada (.docx)", type=["docx"], key="doc_quijada")

if doc_quijada:
    try:
        _doc_check = Document(doc_quijada)
        doc_quijada.seek(0)
        _texto_check = '\n'.join(''.join(p.itertext()) for p in _doc_check.element.body if p.tag == qn('w:p'))
        if 'Paciente:' not in _texto_check:
            st.warning("⚠️ No se encontró la línea 'Paciente:' en este documento. Revísalo antes de continuar, la unión podría fallar o quedar incompleta.")
        if 'Atentamente' not in _texto_check:
            st.warning("⚠️ No se encontró la palabra 'Atentamente' (cierre de firma) en este documento. Revísalo antes de continuar, la unión podría fallar o quedar incompleta.")
    except Exception:
        st.warning("⚠️ No se pudo leer este archivo como un documento Word válido.")

st.markdown("**Datos adicionales (opcional, solo si faltan en el documento):**")
col_fecha, col_medico, col_cedula = st.columns(3)
with col_fecha:
    fecha_estudio = st.date_input("Fecha de Estudio", value=None, format="DD/MM/YYYY")
with col_medico:
    medico_referente = st.text_input("Médico Referente")
with col_cedula:
    cedula_paciente = st.text_input("Cédula del Paciente")

if doc_quijada and st.button("👁️ Generar Vista Previa"):
    try:
        fecha_texto = fecha_estudio.strftime("%d/%m/%Y") if fecha_estudio else None
        merged_bytes, nombre_detectado = unir_informe_renal(
            "plantilla_renal.docx", doc_quijada,
            fecha_estudio=fecha_texto,
            medico_referente=medico_referente,
            cedula=cedula_paciente
        )
        st.session_state["preview_bytes"] = merged_bytes.getvalue()
        st.session_state["preview_nombre"] = nombre_detectado
    except Exception as e:
        st.error(f"Error al unir los documentos: {e}")

if st.session_state.get("preview_bytes"):
    st.markdown("### 👁️ Vista Previa (todavía no se ha guardado)")
    st.caption("Esto es solo texto, para que revises el contenido. El documento Word final sí tendrá el formato y las imágenes completas.")
    with st.container(border=True):
        for linea in extraer_vista_previa(st.session_state["preview_bytes"]):
            st.write(linea)

    col_confirmar, col_descartar = st.columns(2)
    with col_confirmar:
        if st.button("✅ Confirmar y Guardar en Historial", type="primary"):
            timestamp = datetime.datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
            nombre_detectado = st.session_state.get("preview_nombre")
            nombre_base = nombre_detectado.replace(' ', '_') if nombre_detectado else "SinNombre"
            nombre_hist = f"{timestamp}_{nombre_base}.docx"
            guardar_en_historial(nombre_hist, st.session_state["preview_bytes"])
            st.session_state.pop("preview_bytes", None)
            st.session_state.pop("preview_nombre", None)
            st.toast(f"✅ Guardado como '{nombre_hist}'. Lo encuentras más abajo, en Historial de Informes.", icon="✅")
            st.rerun()
    with col_descartar:
        if st.button("🗑️ Descartar"):
            st.session_state.pop("preview_bytes", None)
            st.session_state.pop("preview_nombre", None)
            st.rerun()

st.markdown("---")
st.header("📁 Historial de Informes")
st.markdown("Todos los informes generados (de audio o unidos con el Dr. Quijada) quedan guardados aquí automáticamente. Puedes buscar, renombrar o borrar cualquiera.")

busqueda = st.text_input("🔍 Buscar por nombre de archivo o paciente")

archivos_historial = sorted(os.listdir(CARPETA_HISTORIAL), reverse=True)
if busqueda:
    archivos_historial = [a for a in archivos_historial if busqueda.lower() in a.lower()]

st.caption(f"{len(archivos_historial)} informe(s) encontrados" if busqueda else f"{len(archivos_historial)} informe(s) en total")

if archivos_historial:
    for nombre_archivo in archivos_historial:
        ruta = os.path.join(CARPETA_HISTORIAL, nombre_archivo)
        es_word = nombre_archivo.endswith(".docx")
        icono = "📄" if es_word else "📝"

        fecha_legible = ""
        partes = nombre_archivo.split("_", 2)
        if len(partes) >= 2:
            try:
                fecha_legible = datetime.datetime.strptime(f"{partes[0]}_{partes[1]}", "%Y-%m-%d_%H-%M-%S").strftime("%d/%m/%Y %H:%M")
            except ValueError:
                pass

        with st.container(border=True):
            col_info, col_desc, col_ren, col_del = st.columns([3, 1, 1, 1])

            with col_info:
                st.markdown(f"{icono} **{nombre_archivo}**")
                if fecha_legible:
                    st.caption(fecha_legible)

            with col_desc:
                with open(ruta, "rb") as f:
                    datos_archivo = f.read()
                mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document" if es_word else "text/plain"
                st.download_button("⬇️ Descargar", data=datos_archivo, file_name=nombre_archivo, mime=mime, key=f"desc_{nombre_archivo}")

            with col_ren:
                if st.button("✏️ Renombrar", key=f"ren_btn_{nombre_archivo}"):
                    st.session_state["renombrando"] = nombre_archivo
                    st.session_state.pop("confirmar_borrado", None)

            with col_del:
                if st.button("🗑️ Borrar", key=f"del_btn_{nombre_archivo}"):
                    st.session_state["confirmar_borrado"] = nombre_archivo
                    st.session_state.pop("renombrando", None)

            if st.session_state.get("renombrando") == nombre_archivo:
                extension = os.path.splitext(nombre_archivo)[1]
                nombre_sin_ext = os.path.splitext(nombre_archivo)[0]
                nuevo_nombre = st.text_input("Nuevo nombre (sin extensión):", value=nombre_sin_ext, key=f"input_ren_{nombre_archivo}")
                col_ok, col_cancel = st.columns(2)
                with col_ok:
                    if st.button("✅ Guardar nombre", key=f"ok_ren_{nombre_archivo}"):
                        nueva_ruta = os.path.join(CARPETA_HISTORIAL, nuevo_nombre.strip() + extension)
                        if nuevo_nombre.strip():
                            os.rename(ruta, nueva_ruta)
                        st.session_state.pop("renombrando", None)
                        st.rerun()
                with col_cancel:
                    if st.button("❌ Cancelar", key=f"cancel_ren_{nombre_archivo}"):
                        st.session_state.pop("renombrando", None)
                        st.rerun()

            if st.session_state.get("confirmar_borrado") == nombre_archivo:
                st.warning(f"¿Seguro que quieres borrar '{nombre_archivo}'? Esta acción no se puede deshacer.")
                col_si, col_no = st.columns(2)
                with col_si:
                    if st.button("🗑️ Sí, borrar definitivamente", key=f"si_del_{nombre_archivo}"):
                        os.remove(ruta)
                        st.session_state.pop("confirmar_borrado", None)
                        st.rerun()
                with col_no:
                    if st.button("Cancelar", key=f"no_del_{nombre_archivo}"):
                        st.session_state.pop("confirmar_borrado", None)
                        st.rerun()
else:
    if busqueda:
        st.info("No se encontraron informes que coincidan con la búsqueda.")
    else:
        st.info("Todavía no hay informes generados.")